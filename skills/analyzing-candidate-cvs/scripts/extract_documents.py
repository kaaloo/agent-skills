#!/usr/bin/env python3
"""Extract text from PDF, DOCX, TXT, and Markdown files for CV analysis.

Usage:
  python extract_documents.py /path/to/cv-analysis
  python extract_documents.py /path/to/cv-analysis --out /path/to/_extracted_text
"""
from __future__ import annotations

import argparse
from pathlib import Path

SUPPORTED = {".pdf", ".docx", ".txt", ".md"}


def extract_pdf(path: Path) -> str:
    import pdfplumber

    chunks: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            chunks.append(f"\n--- Page {i} ---\n{text}")
    return "\n".join(chunks)


def extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    chunks: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            chunks.append(text)
    for i, table in enumerate(doc.tables, 1):
        chunks.append(f"\n--- Tableau {i} ---")
        for row in table.rows:
            cells = []
            for cell in row.cells:
                text = cell.text.strip()
                if not cells or text != cells[-1]:
                    cells.append(text)
            if any(cells):
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def extract_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {path.suffix}")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("root", type=Path, help="Hiring campaign root directory")
    parser.add_argument("--out", type=Path, default=None, help="Output directory")
    args = parser.parse_args()

    root = args.root.expanduser().resolve()
    out = (args.out or (root / "_extracted_text")).expanduser().resolve()
    out.mkdir(parents=True, exist_ok=True)

    ignored_dirs = {"_extracted_text", "_reports", ".git", "node_modules"}
    files = [
        p for p in root.rglob("*")
        if p.is_file()
        and p.suffix.lower() in SUPPORTED
        and out not in p.parents
        and not (ignored_dirs & set(p.parts))
    ]

    for path in sorted(files):
        rel = path.relative_to(root)
        target = out / (rel.as_posix().replace("/", "__") + ".txt")
        try:
            text = extract_file(path)
            target.write_text(text, encoding="utf-8")
            print(f"OK {rel} -> {target.name} ({len(text)} chars)")
        except Exception as exc:
            print(f"ERROR {rel}: {exc}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
